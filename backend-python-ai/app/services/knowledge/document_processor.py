"""
文档处理服务 - 基于LangChain简化实现
处理文档解析、文本分割
"""

import os
import uuid
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from langchain.schema import Document
try:
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain_community.document_loaders import PyPDFLoader, TextLoader, UnstructuredWordDocumentLoader
except ImportError:
    # 兼容旧版本
    from langchain.text_splitter import RecursiveCharacterTextSplitter

from app.utils.logger import LoggerMixin


class DocumentProcessor(LoggerMixin):
    """
    文档处理器
    基于LangChain的文档加载和分割
    """
    
    def __init__(self, upload_dir: str = "./data/uploads"):
        """初始化文档处理器"""
        super().__init__()
        
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        
        # 初始化文本分割器
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", "。", "！", "？", ". ", "! ", "? ", " ", ""],
            length_function=len,
        )
        
        self.log_info("文档处理器初始化完成")
    
    async def save_uploaded_file(
        self, 
        file_content: bytes, 
        filename: str,
        knowledge_base_id: str
    ) -> Dict[str, Any]:
        """
        保存上传的文件
        
        Args:
            file_content: 文件内容
            filename: 文件名
            knowledge_base_id: 知识库ID
            
        Returns:
            Dict: 文件信息
        """
        try:
            # 生成文件ID
            file_id = f"file_{uuid.uuid4().hex[:8]}"
            
            # 创建知识库目录
            kb_dir = self.upload_dir / knowledge_base_id
            kb_dir.mkdir(parents=True, exist_ok=True)
            
            # 保存文件
            file_ext = Path(filename).suffix.lower()
            saved_filename = f"{file_id}{file_ext}"
            file_path = kb_dir / saved_filename
            
            with open(file_path, "wb") as f:
                f.write(file_content)
            
            file_info = {
                "file_id": file_id,
                "file_name": filename,
                "file_path": str(file_path),
                "file_size": len(file_content),
                "file_type": file_ext.lstrip("."),
                "knowledge_base_id": knowledge_base_id,
                "upload_time": datetime.now().isoformat(),
            }
            
            self.log_info(f"文件保存成功: {filename} -> {file_path}")
            return file_info
            
        except Exception as e:
            self.log_error(f"保存文件失败: {filename}, 错误: {str(e)}", error=e)
            raise
    
    async def parse_document(
        self, 
        file_path: str, 
        file_type: Optional[str] = None
    ) -> List[Document]:
        """
        解析文档
        
        Args:
            file_path: 文件路径
            file_type: 文件类型（可选，自动检测）
            
        Returns:
            List[Document]: LangChain文档列表
        """
        try:
            path = Path(file_path)
            
            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 自动检测文件类型
            if file_type is None:
                file_type = path.suffix.lower().lstrip(".")
            
            self.log_info(f"开始解析文档: {path.name}, 类型: {file_type}")
            
            # 根据文件类型选择加载器
            if file_type == "pdf":
                documents = await self._load_pdf(path)
            elif file_type in ["docx", "doc"]:
                documents = await self._load_word(path)
            elif file_type in ["txt", "md", "markdown"]:
                documents = await self._load_text(path)
            else:
                # 默认按文本处理
                documents = await self._load_text(path)
            
            self.log_info(f"文档解析完成: {path.name}, 共 {len(documents)} 页/段落")
            return documents
            
        except Exception as e:
            self.log_error(f"解析文档失败: {file_path}, 错误: {str(e)}", error=e)
            raise
    
    async def _load_pdf(self, file_path: Path) -> List[Document]:
        """加载PDF文件"""
        try:
            try:
                from langchain_community.document_loaders import PyPDFLoader
            except ImportError:
                from langchain.document_loaders import PyPDFLoader
            loader = PyPDFLoader(str(file_path))
            return loader.load()
        except ImportError:
            self.log_error("PyPDF2未安装，请运行: pip install pypdf2")
            raise
    
    async def _load_word(self, file_path: Path) -> List[Document]:
        """加载Word文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(str(file_path))
            full_text = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    full_text.append(text)
            
            # 合并所有段落
            content = "\n".join(full_text)
            
            return [Document(
                page_content=content,
                metadata={"source": str(file_path), "file_type": "docx"}
            )]
        except ImportError:
            self.log_error("python-docx未安装，请运行: pip install python-docx")
            raise
        except Exception as e:
            self.log_error(f"读取Word文档失败: {str(e)}")
            raise
    
    async def _load_text(self, file_path: Path) -> List[Document]:
        """加载文本文件"""
        try:
            try:
                from langchain_community.document_loaders import TextLoader
            except ImportError:
                from langchain.document_loaders import TextLoader
            loader = TextLoader(str(file_path), encoding="utf-8")
            return loader.load()
        except Exception as e:
            # 如果编码错误，尝试其他编码
            try:
                try:
                    from langchain_community.document_loaders import TextLoader
                except ImportError:
                    from langchain.document_loaders import TextLoader
                loader = TextLoader(str(file_path), encoding="gbk")
                return loader.load()
            except:
                raise e
    
    async def split_documents(
        self, 
        documents: List[Document],
        chunk_size: int = 1000,
        chunk_overlap: int = 200
    ) -> List[Document]:
        """
        分割文档
        
        Args:
            documents: 文档列表
            chunk_size: 分块大小
            chunk_overlap: 重叠大小
            
        Returns:
            List[Document]: 分割后的文档块
        """
        try:
            # 更新分割器参数
            self.text_splitter.chunk_size = chunk_size
            self.text_splitter.chunk_overlap = chunk_overlap
            
            self.log_info(f"文档分割参数: chunk_size={chunk_size}, chunk_overlap={chunk_overlap}")
            
            # 分割文档
            chunks = self.text_splitter.split_documents(documents)
            
            self.log_info(f"文档分割完成: {len(documents)} -> {len(chunks)} 块")
            return chunks
            
        except Exception as e:
            self.log_error(f"分割文档失败: {str(e)}", error=e)
            raise
    
    async def process_file(
        self,
        file_content: bytes,
        filename: str,
        knowledge_base_id: str,
        config: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        完整处理流程：保存 -> 解析 -> 分割
        
        Args:
            file_content: 文件内容
            filename: 文件名
            knowledge_base_id: 知识库ID
            config: 处理配置
            
        Returns:
            Dict: 处理结果
        """
        config = config or {}
        
        try:
            # 1. 保存文件
            file_info = await self.save_uploaded_file(
                file_content, filename, knowledge_base_id
            )
            
            # 2. 解析文档
            documents = await self.parse_document(
                file_info["file_path"], 
                file_info["file_type"]
            )
            
            # 3. 分割文档
            chunk_size = config.get("knowledge_length", 1000)
            chunk_overlap = config.get("overlap_length", 200)
            
            chunks = await self.split_documents(
                documents, chunk_size, chunk_overlap
            )
            
            # 4. 构建结果
            result = {
                "file_id": file_info["file_id"],
                "file_name": filename,
                "file_size": file_info["file_size"],
                "file_type": file_info["file_type"],
                "knowledge_base_id": knowledge_base_id,
                "total_chunks": len(chunks),
                "chunks": [
                    {
                        "content": chunk.page_content,
                        "metadata": chunk.metadata,
                        "word_count": len(chunk.page_content),
                    }
                    for chunk in chunks
                ],
                "parse_status": "completed",
            }
            
            self.log_info(
                f"文件处理完成: {filename}, "
                f"生成 {len(chunks)} 个知识块"
            )
            
            return result
            
        except Exception as e:
            self.log_error(f"处理文件失败: {filename}, 错误: {str(e)}", error=e)
            return {
                "file_name": filename,
                "knowledge_base_id": knowledge_base_id,
                "parse_status": "failed",
                "error": str(e),
            }
    
    def get_file_path(self, file_id: str, knowledge_base_id: str) -> Optional[Path]:
        """获取文件路径"""
        kb_dir = self.upload_dir / knowledge_base_id
        
        # 查找匹配的文件
        for file_path in kb_dir.glob(f"{file_id}*"):
            return file_path
        
        return None
    
    async def delete_file(self, file_id: str, knowledge_base_id: str) -> bool:
        """删除文件"""
        try:
            file_path = self.get_file_path(file_id, knowledge_base_id)
            if file_path and file_path.exists():
                file_path.unlink()
                self.log_info(f"文件删除成功: {file_id}")
                return True
            return False
        except Exception as e:
            self.log_error(f"删除文件失败: {file_id}, 错误: {str(e)}", error=e)
            return False


# 创建全局实例
document_processor = DocumentProcessor()
